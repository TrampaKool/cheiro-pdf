import streamlit as st
from google import genai
import time
from dotenv import load_dotenv
import os
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
from reportlab.pdfbase.pdfmetrics import stringWidth
import random
from natsort import natsorted
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MAX_WIDTH = 500  # Total pdf width in points

load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=API_KEY)

# --- UI LAYOUT ---
st.set_page_config(page_title="CheiroPDF")
st.title("Handwritten Greek to PDF")
st.write("Upload your notes, and convert them to an editable PDF.")

if not API_KEY:
    st.error("API Key not found! Check your .env file.")
    st.stop()

# Register a font that supports Greek (Must have the .ttf file in your folder)
try:
    pdfmetrics.registerFont(TTFont('GreekFont', 'DejaVuSans.ttf'))
    FONT_NAME = 'GreekFont'
except:
    st.warning("Greek font not found. Using default (may show boxes in PDF).")
    FONT_NAME = 'Helvetica'

uploaded_files = st.file_uploader(
    "Choose images...", type=["jpg", "jpeg", "png"], accept_multiple_files=True
)
sorted_files = natsorted(uploaded_files, key=lambda x: x.name)
missed_pages = []

if sorted_files:
    st.info(f"Loaded {len(sorted_files)} images.")

     # --- OUTPUT FORMAT SELECTOR ---
    use_docx = st.toggle("Export as DOCX", value=False)
    output_format = "DOCX" if use_docx else "PDF"
    
    if st.button("Start Transcription"):
        all_text = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, file in enumerate(sorted_files):
            # Update UI
            percent = (i + 1) / len(sorted_files)
            progress_bar.progress(percent)
            status_text.text(f"Processing image {i+1} of {len(sorted_files)}...")

            # Convert uploaded file to bytes for the API
            img_bytes = file.read()
            
            # --- API CALL WITH RETRY & BACKOFF ---
            max_retries = 7
            retry_count = 0
            success = False
            response_text = ""

            while retry_count < max_retries and not success:
                try:
                    response = client.models.generate_content(
                        model="gemini-3.1-flash-lite-preview",
                        contents=[
                            "Transcribe the handwritten Greek text exactly. Preserve line breaks. Do not translate. Return only the transcription.",
                            {"inline_data": {"data": img_bytes, "mime_type": "image/jpeg"}}
                        ]
                    )
                    response_text = response.text
                    success = True
                    break # Exit retry loop on success
                
                except Exception as e:
                    error_msg = str(e)
                    retry_count += 1
                    print(error_msg)
                    # Exponential backoff: 2, 4, 8, 16, 32 seconds + jitter

                    wait_time = min((2 ** retry_count) + random.uniform(0, 1), 8)
                    status_text.warning(f"Rate limited. Retrying in {wait_time:.1f}s... (Attempt {retry_count}/{max_retries})")
                    time.sleep(wait_time)

            if success:
                all_text.append(response.text or "")
            else:
                missed_pages.append(file.name)
                all_text.append(f"[Error: Could not transcribe image {i+1} Name: {file.name}]")

        if (len(missed_pages) > 0):
            status_text.warning(f"Pages that failed to process: {",".join(missed_pages)}")
        else:
            status_text.success("All images processed!")
        
        
        if output_format == "PDF":
            # --- PDF GENERATION ---
            pdf_buffer = io.BytesIO()
            c = canvas.Canvas(pdf_buffer)
            
            # Simple PDF assembly
            y_position = 800
            c.setFont(FONT_NAME, 12)
            
            for text in all_text:
                lines = text.split('\n')
                for line in lines:
                    current_line = ""
                    words = line.split(' ')
        
                    for word in words:
                        # Check if adding the next word exceeds our MAX_WIDTH
                        test_line = f"{current_line} {word}".strip()
                        width = stringWidth(test_line, FONT_NAME, 12)
                        
                        if width < MAX_WIDTH:
                            current_line = test_line
                        else:
                            # The line is full! Print it and start a new one
                            c.drawString(50, y_position, current_line)
                            y_position -= 15
                            current_line = word
                            
                            # Check for page break
                            if y_position < 50:
                                c.showPage()
                                c.setFont(FONT_NAME, 12)
                                y_position = 800
                                
                    # Draw the final leftover piece of the line
                    c.drawString(50, y_position, current_line)
                    y_position -= 15

                c.showPage()
                c.setFont(FONT_NAME, 12)
                y_position = 800

            c.save()
            pdf_buffer.seek(0)

            # --- DOWNLOAD BUTTON ---
            st.download_button(
                label="📩 Download Transcribed PDF",
                data=pdf_buffer,
                file_name="transcribed_greek.pdf",
                mime="application/pdf"
            )

        # -------------------------------------------------------
        # --- DOCX GENERATION -----------------------------------
        # -------------------------------------------------------
        else:
            doc = DocxDocument()
 
            # Set default font to one that supports Greek (DejaVu Sans or fallback)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'DejaVu Sans'
            font.size = Pt(12)
 
            # Helper: add a page break paragraph between pages
            def add_page_break(document):
                paragraph = document.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._r.append(br)
 
            for page_index, text in enumerate(all_text):

                lines = text.split('\n')
 
                for line_index, line in enumerate(lines):
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.font.name = 'DejaVu Sans'
                    run.font.size = Pt(12)
 
                # After each image's text block, insert a page break
                # (except after the very last page to avoid a trailing blank page)
                if page_index < len(all_text) - 1:
                    add_page_break(doc)
 
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
 
            # --- DOWNLOAD BUTTON ---
            st.download_button(
                label="📩 Download Transcribed DOCX",
                data=docx_buffer,
                file_name="transcribed_greek.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )